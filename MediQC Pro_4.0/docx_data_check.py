import os
import re
import pandas as pd
from docx import Document
import json
from datetime import datetime
import logging

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocxDataExtractor:
    """
    从Word文档中提取医疗记录数据的类
    """
    def __init__(self):
        """初始化提取器"""
        # 定义常见的字段模式
        self.field_patterns = {
            '姓名': r'姓名[_\s]*([^\s_]+)[_\s]*',
            '性别': r'性别[_\s]*([男女12])[_\s]*',
            '年龄': r'年龄[_\s]*(\d+)[_\s]*',
            '出生日期': r'出生日期[_\s]*(\d{4})[年_\s]+(\d{1,2})[月_\s]+(\d{1,2})[日_\s]*',
            '国籍': r'国籍[_\s]*([^\s_]+)[_\s]*',
            '出生地址': r'出生地址[_\s]*([^\n]+)[_\s]*',
            '民族': r'民族[_\s]*([^\s_]+)[_\s]*',
            '证件类别': r'证件类别[_\s]*([^\s_]+)[_\s]*',
            '证件号码': r'(?:证件号码|身份证号)[_\s]*([0-9Xx]{15,18}|[^\s_]+)[_\s]*',
            '入院途径': r'入院途径[_\s]*([^\s_]+)[_\s]*',
            '入院时间': r'(?:入院时间|入院日期)[_\s]*(\d{4})[-年_\s]+(\d{1,2})[-月_\s]+(\d{1,2})(?:[-日_\s]+(\d{1,2}):(\d{1,2}))?[_\s]*',
            '入院科别': r'(?:入院科别|科室)[_\s]*([^\s_]+)[_\s]*',
            '入院病房': r'(?:入院病房|病房|床号)[_\s]*([^\s_]+)[_\s]*',
            '转科科别': r'转科科别[_\s]*([^\s_]+)[_\s]*',
            '出院时间': r'(?:出院时间|出院日期)[_\s]*(\d{4})[-年_\s]+(\d{1,2})[-月_\s]+(\d{1,2})(?:[-日_\s]+(\d{1,2}):(\d{1,2}))?[_\s]*',
            '出院科别': r'出院科别[_\s]*([^\s_]+)[_\s]*',
            '出院病房': r'出院病房[_\s]*([^\s_]+)[_\s]*',
            '实际住院': r'(?:实际住院|住院天数)[_\s]*(\d+)[天_\s]*',
            '门诊诊断名称': r'(?:门诊诊断名称|门急诊诊断名称)[_\s]*([^\n]+)[_\s]*',
            '门诊诊断编码': r'(?:门诊诊断编码|门急诊诊断编码)[_\s]*([^\n]+)[_\s]*',
            '入院诊断名称': r'入院诊断(?:名称)?[_\s]*([^\n]+)[_\s]*',
            '入院诊断编码': r'入院诊断编码[_\s]*([^\n]+)[_\s]*',
            '主要诊断': r'主要诊断[_\s]*([^\n]+)[_\s]*',
            '其他诊断': r'其他诊断[_\s]*([^\n]+)[_\s]*',
            '损伤中毒的外部原因': r'(?:损伤中毒的外部原因|损伤、中毒的外部原因)[_\s]*([^\n]+)[_\s]*',
            '病理诊断名称': r'病理诊断(?:名称)?[_\s]*([^\n]+)[_\s]*',
            '病理诊断编码': r'病理诊断编码[_\s]*([^\n]+)[_\s]*',
            '病理描述': r'病理描述[_\s]*([^\n]+)[_\s]*',
            '病理号': r'病理号[_\s]*([^\s_]+)[_\s]*',
            '是否有药物过敏': r'是否有药物过敏[_\s]*([是否有无YyNn01是的不是]+)[_\s]*',
            '过敏药物': r'(?:过敏药物|过敏药物源)[_\s]*([^\n]*)[_\s]*',
            '是否进行了死亡患者尸检': r'是否进行了死亡患者尸检[_\s]*([是否有无YyNn01是的不是]+)[_\s]*',
        }
        
        # 字段名称映射，用于处理不同表述的字段名
        self.field_name_mapping = {
            '姓名': '姓名',
            '性别': '性别',
            '年龄': '年龄',
            '出生日期': '出生日期',
            '国籍': '国籍',
            '出生地': '出生地址',
            '出生地址': '出生地址',
            '民族': '民族',
            '证件类别': '证件类别',
            '证件号码': '证件号码',
            '身份证号': '证件号码',
            '职业': '职业',
            '入院途径': '入院途径',
            '入院时间': '入院时间',
            '入院日期': '入院时间',
            '入院科别': '入院科别',
            '科室': '入院科别',
            '入院病房': '入院病房',
            '病房': '入院病房',
            '床号': '入院病房',
            '转科科别': '转科科别',
            '出院时间': '出院时间',
            '出院日期': '出院时间',
            '出院科别': '出院科别',
            '出院病房': '出院病房',
            '实际住院': '实际住院',
            '住院天数': '实际住院',
            '门诊诊断名称': '门诊诊断名称',
            '门急诊诊断名称': '门诊诊断名称',
            '门诊诊断编码': '门诊诊断编码',
            '门急诊诊断编码': '门诊诊断编码',
            '入院诊断名称': '入院诊断名称',
            '入院诊断': '入院诊断名称',
            '入院诊断编码': '入院诊断编码',
            '主要诊断': '主要诊断',
            '主要诊断编码': '主要诊断编码',
            '其他诊断': '其他诊断',
            '其他诊断编码': '其他诊断编码',
            '损伤中毒的外部原因': '损伤中毒的外部原因',
            '损伤、中毒的外部原因': '损伤中毒的外部原因',
            '病理诊断名称': '病理诊断名称',
            '病理诊断': '病理诊断名称',
            '病理诊断编码': '病理诊断编码',
            '病理描述': '病理描述',
            '病理号': '病理号',
            '是否有药物过敏': '是否有药物过敏',
            '过敏药物': '过敏药物',
            '过敏药物源': '过敏药物',
            '是否进行了死亡患者尸检': '是否进行了死亡患者尸检'
        }
        
        # 性别映射
        self.gender_mapping = {
            '1': '男',
            '2': '女',
            '男': '男',
            '女': '女',
            'M': '男',
            'F': '女',
            'm': '男',
            'f': '女'
        }
        
        # 是/否映射
        self.yes_no_mapping = {
            '1': '是',
            '0': '否',
            'Y': '是',
            'N': '否',
            'y': '是',
            'n': '否',
            '是': '是',
            '否': '否',
            '有': '是',
            '无': '否',
            '是的': '是',
            '不是': '否'
        }
        
        # 入院途径映射
        self.admission_path_mapping = {
            '1': '门诊',
            '2': '急诊',
            '3': '其他医疗机构转入',
            '9': '其他',
            '门诊': '门诊',
            '急诊': '急诊',
            '转入': '其他医疗机构转入',
            '其他医疗机构转入': '其他医疗机构转入',
            '其他': '其他'
        }
    
    def extract_data(self, text):
        """
        从文本中提取数据，使用混合提取策略
        
        Args:
            text (str): 要处理的文本
            
        Returns:
            dict: 提取的数据字段及其值
        """
        # 首先尝试使用分号提取
        data = self.extract_data_by_semicolon(text)
        
        # 如果分号提取的结果不完整，再使用正则表达式提取
        if len(data) < 10:  # 假设至少应该有10个字段
            regex_data = self.extract_data_by_regex(text)
            
            # 合并两种提取方法的结果，优先使用分号提取的结果
            for key, value in regex_data.items():
                if key not in data or not data[key]:
                    data[key] = value
        
        # 确保在任何情况下都处理混合字段
        self._separate_mixed_fields(data)
        
        # 针对特定错误模式进行修复
        self._fix_known_errors(data)
        
        # 清理数据
        cleaned_data = self.clean_data(data)
        
        return cleaned_data
    
    def extract_data_by_semicolon(self, text):
        """
        使用分号分隔符提取数据
        
        Args:
            text (str): 要处理的文本
            
        Returns:
            dict: 提取的数据字段及其值
        """
        data = {}
        
        # 统一处理中英文分号
        text = text.replace(';', '；')
        
        # 使用正则表达式查找所有的字段-值对
        # 模式: 字段名[：:=]值[；\n]
        pattern = r'([\u4e00-\u9fa5a-zA-Z0-9]+)[：:=](.*?)(?=；|$)'
        
        # 查找包含分号的行
        lines = text.split('\n')
        for line in lines:
            if '；' in line:
                # 预处理：将空字段标记出来
                # 例如：将"姓名；性别：2"替换为"姓名：；性别：2"
                for field_name in self.field_name_mapping.keys():
                    line = re.sub(f"({field_name})；", r"\1：；", line)
                
                # 分割行，获取所有键值对
                pairs = line.split('；')
                for pair in pairs:
                    pair = pair.strip()
                    if not pair:
                        continue
                        
                    # 查找键值对
                    if ':' in pair or '：' in pair or '=' in pair:
                        # 处理包含冒号或等号的情况
                        if ':' in pair:
                            parts = pair.split(':', 1)
                        elif '：' in pair:
                            parts = pair.split('：', 1)
                        elif '=' in pair:
                            parts = pair.split('=', 1)
                            
                        if len(parts) == 2:
                            key = parts[0].strip()
                            value = parts[1].strip()
                            
                            # 如果值为空或只包含下划线，则视为空值
                            if not value or re.match(r'^_+$', value):
                                value = ""
                                
                            # 检查值是否包含其他字段名（可能是错误合并的）
                            for next_field in self.field_name_mapping.keys():
                                if value and value.startswith(next_field):
                                    value = ""
                                    break
                                    
                            # 映射字段名称
                            if key in self.field_name_mapping:
                                standardized_key = self.field_name_mapping[key]
                                
                                # 处理特殊字段
                                if value:  # 只有在值不为空时才处理
                                    if standardized_key == '性别':
                                        value = self.gender_mapping.get(value, value)
                                    elif standardized_key == '入院途径':
                                        # 处理入院途径代码
                                        code_match = re.match(r'^(\d+)', value)
                                        if code_match:
                                            code = code_match.group(1)
                                            if code in self.admission_path_mapping:
                                                value = self.admission_path_mapping[code]
                                        else:
                                            value = self.admission_path_mapping.get(value, value)
                                    elif standardized_key in ['是否有药物过敏', '是否进行了死亡患者尸检']:
                                        value = self.yes_no_mapping.get(value, value)
                                    elif standardized_key in ['出生日期', '入院时间', '出院时间']:
                                        # 尝试解析日期
                                        value = self.parse_date_time(value, standardized_key)
                                
                                data[standardized_key] = value
                    else:
                        # 处理只有字段名没有值的情况（空字段）
                        key = pair.strip()
                        if key in self.field_name_mapping:
                            data[self.field_name_mapping[key]] = ""
        
        # 使用正则表达式查找可能的键值对
        # 这是一个备用方法，用于处理可能的特殊格式
        matches = re.findall(pattern, text)
        for key, value in matches:
            key = key.strip()
            value = value.strip()
            
            # 如果值为空或只包含下划线，则视为空值
            if not value or re.match(r'^_+$', value):
                value = ""
                
            # 映射字段名称
            if key in self.field_name_mapping:
                standardized_key = self.field_name_mapping[key]
                
                # 如果数据中已经有这个字段且不为空，则不覆盖
                if standardized_key not in data or not data[standardized_key]:
                    # 处理特殊字段
                    if value:  # 只有在值不为空时才处理
                        if standardized_key == '性别':
                            value = self.gender_mapping.get(value, value)
                        elif standardized_key == '入院途径':
                            # 处理入院途径代码
                            code_match = re.match(r'^(\d+)', value)
                            if code_match:
                                code = code_match.group(1)
                                if code in self.admission_path_mapping:
                                    value = self.admission_path_mapping[code]
                            else:
                                value = self.admission_path_mapping.get(value, value)
                        elif standardized_key in ['是否有药物过敏', '是否进行了死亡患者尸检']:
                            value = self.yes_no_mapping.get(value, value)
                        elif standardized_key in ['出生日期', '入院时间', '出院时间']:
                            # 尝试解析日期
                            value = self.parse_date_time(value, standardized_key)
                    
                    data[standardized_key] = value
        
        # 处理表格数据：根据图片所示的问题，特别处理可能混合在一起的字段
        self._separate_mixed_fields(data)
        
        # 针对特定字段的专门处理
        self._process_special_fields(text, data)
        
        # 清理数据，确保没有错误的合并字段
        self._clean_merged_fields(data)
        
        # 最后清理数据中的下划线
        for key in data:
            if isinstance(data[key], str) and re.match(r'^_+$', data[key]):
                data[key] = ""
        
        return data
    
    def _separate_mixed_fields(self, data):
        """
        分离混合在一起的字段，如入院诊断名称和编码、病理诊断名称和编码等
        
        Args:
            data (dict): 提取的数据字典
        """
        # 处理出生地址、籍贯和民族
        if '出生地址' in data:
            value = data['出生地址']
            # 检查是否包含民族信息
            ethnicity_match = re.search(r'民族[:：\s]*([\u4e00-\u9fa5]+族?)', value)
            if ethnicity_match:
                ethnicity = ethnicity_match.group(1)
                # 更新出生地址，去除民族信息
                data['出生地址'] = value.replace(ethnicity_match.group(0), '').strip()
                # 如果没有民族字段，添加它
                if '民族' not in data or not data['民族']:
                    data['民族'] = ethnicity
            
            # 检查是否包含籍贯信息
            origin_match = re.search(r'籍贯[:：\s]*([\u4e00-\u9fa5]+)', value)
            if origin_match:
                # 更新出生地址，去除籍贯信息
                data['出生地址'] = value.replace(origin_match.group(0), '').strip()
        
        # 处理入院诊断名称和编码
        if '入院诊断名称' in data:
            value = data['入院诊断名称']
            # 检查是否包含编码（通常是字母+数字的组合）
            code_match = re.search(r'([A-Z]\d+\.\d+|[A-Z]\d+)', value)
            if code_match:
                code = code_match.group(0)
                # 尝试分离名称和编码
                name_parts = value.split(code)
                if len(name_parts) > 0:
                    # 更新名称，去除可能的分隔符
                    data['入院诊断名称'] = re.sub(r'[;；,，、]?\s*$', '', name_parts[0].strip())
                    # 如果没有编码字段，添加它
                    if '入院诊断编码' not in data or not data['入院诊断编码']:
                        data['入院诊断编码'] = code
        
        # 处理主要诊断和编码
        if '主要诊断' in data:
            value = data['主要诊断']
            code_match = re.search(r'([A-Z]\d+\.\d+|[A-Z]\d+)', value)
            if code_match:
                code = code_match.group(0)
                name_parts = value.split(code)
                if len(name_parts) > 0:
                    data['主要诊断'] = re.sub(r'[;；,，、]?\s*$', '', name_parts[0].strip())
                    # 主要诊断编码可能没有单独的字段，但我们可以添加到其他诊断编码中
                    if '主要诊断编码' not in data:
                        data['主要诊断编码'] = code
        
        # 处理病理诊断名称和编码
        if '病理诊断名称' in data:
            value = data['病理诊断名称']
            code_match = re.search(r'([A-Z]\d+\.\d+|[A-Z]\d+)', value)
            if code_match:
                code = code_match.group(0)
                name_parts = value.split(code)
                if len(name_parts) > 0:
                    data['病理诊断名称'] = re.sub(r'[;；,，、]?\s*$', '', name_parts[0].strip())
                    if '病理诊断编码' not in data or not data['病理诊断编码']:
                        data['病理诊断编码'] = code
        
        # 处理病理描述和病理号
        if '病理描述' in data:
            value = data['病理描述']
            # 病理号通常是数字或特定格式的编号
            path_num_match = re.search(r'病理号[：:]\s*([A-Z0-9-]+)', value)
            if path_num_match:
                path_num = path_num_match.group(1)
                # 更新病理描述，去除病理号部分
                data['病理描述'] = value.replace(path_num_match.group(0), '').strip()
                if '病理号' not in data or not data['病理号']:
                    data['病理号'] = path_num
            # 另一种情况：病理号可能在末尾
            else:
                path_num_match = re.search(r'([A-Z0-9-]+)\s*$', value)
                if path_num_match and len(path_num_match.group(1)) >= 5:  # 假设病理号至少5个字符
                    path_num = path_num_match.group(1)
                    data['病理描述'] = value[:path_num_match.start()].strip()
                    if '病理号' not in data or not data['病理号']:
                        data['病理号'] = path_num
    
    def _process_special_fields(self, text, data):
        """
        针对特定字段(入院诊断名称、损伤中毒的外部原因、病理诊断名称、病理描述、过敏药物)的专门处理
        
        Args:
            text (str): 原始文本
            data (dict): 已提取的数据字典
        """
        # 定义特殊字段的提取模式
        special_field_patterns = {
            '出生地址': [
                r'出生地(?:址)?[:：]\s*([^;；\n]+)',
                r'出生地(?:址)?[_\s]+([^;；\n]+)',
                r'出生地(?:址)?[_\s]*(.+?)(?=民族|籍贯|证件|$)'
            ],
            '入院途径': [
                r'入院途径[:：]\s*([^;；\n]+)',
                r'入院途径[_\s]+([^;；\n]+)',
                r'入院途径[_\s]*(.+?)(?=入院时间|入院日期|$)'
            ],
            '入院诊断名称': [
                r'入院诊断(?:名称)?[:：]\s*([^;；\n]+)',
                r'入院诊断(?:名称)?[_\s]+([^;；\n]+)',
                r'入院诊断(?:名称)?[_\s]*(.+?)(?=入院诊断编码|主要诊断|其他诊断|$)'
            ],
            '损伤中毒的外部原因': [
                r'(?:损伤中毒的外部原因|损伤、中毒的外部原因)[:：]\s*([^;；\n]+)',
                r'(?:损伤中毒的外部原因|损伤、中毒的外部原因)[_\s]+([^;；\n]+)',
                r'(?:损伤中毒的外部原因|损伤、中毒的外部原因)[_\s]*(.+?)(?=病理诊断|$)'
            ],
            '病理诊断名称': [
                r'病理诊断(?:名称)?[:：]\s*([^;；\n]+)',
                r'病理诊断(?:名称)?[_\s]+([^;；\n]+)',
                r'病理诊断(?:名称)?[_\s]*(.+?)(?=病理诊断编码|病理描述|$)'
            ],
            '病理描述': [
                r'病理描述[:：]\s*([^;；\n]+)',
                r'病理描述[_\s]+([^;；\n]+)',
                r'病理描述[_\s]*(.+?)(?=病理号|是否有药物过敏|$)'
            ],
            '过敏药物': [
                r'(?:过敏药物|过敏药物源)[:：]\s*([^;；\n]+)',
                r'(?:过敏药物|过敏药物源)[_\s]+([^;；\n]+)',
                r'(?:过敏药物|过敏药物源)[_\s]*(.+?)(?=是否进行了死亡患者尸检|$)'
            ]
        }
        
        # 处理每个特殊字段
        for field, patterns in special_field_patterns.items():
            # 如果字段已经有值且不是空值，则跳过
            if field in data and data[field] and data[field].strip():
                continue
                
            # 尝试每个模式
            for pattern in patterns:
                match = re.search(pattern, text, re.DOTALL)
                if match:
                    value = match.group(1).strip()
                    # 清理值
                    value = re.sub(r'^\s*[:：]\s*', '', value)  # 移除开头可能的冒号
                    value = value.strip()
                    
                    # 处理特殊情况
                    if field == '出生地址':
                        # 处理出生地址中可能包含的民族和籍贯信息
                        ethnicity_match = re.search(r'民族[:：\s]*([\u4e00-\u9fa5]+族?)', value)
                        if ethnicity_match:
                            ethnicity = ethnicity_match.group(1)
                            value = value.replace(ethnicity_match.group(0), '').strip()
                            if '民族' not in data or not data['民族']:
                                data['民族'] = ethnicity
                        
                        origin_match = re.search(r'籍贯[:：\s]*([\u4e00-\u9fa5]+)', value)
                        if origin_match:
                            value = value.replace(origin_match.group(0), '').strip()
                    
                    elif field == '入院途径':
                        # 处理可能包含说明的入院途径，如"2 1.急诊，2门诊，3其他医疗机构转入，9其他"
                        # 先尝试提取开头的数字作为代码
                        code_match = re.match(r'^(\d+)', value)
                        if code_match:
                            code = code_match.group(1)
                            # 使用代码映射
                            if code in self.admission_path_mapping:
                                value = self.admission_path_mapping[code]
                            # 如果代码不在映射中，保留原值
                        else:
                            # 如果没有找到代码，尝试使用整个值进行映射
                            value = self.admission_path_mapping.get(value, value)
                    
                    elif field == '入院诊断名称':
                        # 移除可能包含的编码
                        code_match = re.search(r'([A-Z]\d+\.\d+|[A-Z]\d+)', value)
                        if code_match:
                            code = code_match.group(0)
                            value = value.replace(code, '').strip()
                            if '入院诊断编码' not in data or not data['入院诊断编码']:
                                data['入院诊断编码'] = code
                    
                    elif field == '病理诊断名称':
                        # 移除可能包含的编码
                        code_match = re.search(r'([A-Z]\d+\.\d+|[A-Z]\d+)', value)
                        if code_match:
                            code = code_match.group(0)
                            value = value.replace(code, '').strip()
                            if '病理诊断编码' not in data or not data['病理诊断编码']:
                                data['病理诊断编码'] = code
                    
                    elif field == '病理描述':
                        # 移除可能包含的病理号
                        path_num_match = re.search(r'病理号[：:]\s*([A-Z0-9-]+)', value)
                        if path_num_match:
                            path_num = path_num_match.group(1)
                            value = value.replace(path_num_match.group(0), '').strip()
                            if '病理号' not in data or not data['病理号']:
                                data['病理号'] = path_num
                    
                    # 更新数据字典
                    data[field] = value
                    break
    
    def extract_data_by_regex(self, text):
        """
        使用正则表达式提取数据
        
        Args:
            text (str): 要处理的文本
            
        Returns:
            dict: 提取的数据字段及其值
        """
        data = {}
        
        # 提取数据
        for field, pattern in self.field_patterns.items():
            match = re.search(pattern, text)
            if match:
                if field == '性别':
                    gender_code = match.group(1)
                    data[field] = self.gender_mapping.get(gender_code, gender_code)
                elif field in ['是否有药物过敏', '是否进行了死亡患者尸检']:
                    # 是否类型字段处理
                    yes_no_value = match.group(1)
                    data[field] = self.yes_no_mapping.get(yes_no_value, yes_no_value)
                elif field in ['出生日期']:
                    # 日期字段需要特殊处理
                    if len(match.groups()) >= 3:
                        year = match.group(1)
                        month = match.group(2).zfill(2)  # 补零确保两位数
                        day = match.group(3).zfill(2)    # 补零确保两位数
                        data[field] = f"{year}.{month}.{day}"
                elif field in ['入院时间', '出院时间']:
                    # 日期时间字段需要特殊处理
                    if len(match.groups()) >= 3:
                        year = match.group(1)
                        month = match.group(2).zfill(2)  # 补零确保两位数
                        day = match.group(3).zfill(2)    # 补零确保两位数
                        
                        # 如果有时分信息
                        if len(match.groups()) >= 5 and match.group(4) and match.group(5):
                            hour = match.group(4).zfill(2)
                            minute = match.group(5).zfill(2)
                            data[field] = f"{year}.{month}.{day} {hour}:{minute}"
                        else:
                            data[field] = f"{year}.{month}.{day}"
                else:
                    data[field] = match.group(1).strip()
        
        return data
    
    def parse_date_time(self, value, field):
        """
        解析日期时间字符串
        
        Args:
            value (str): 日期时间字符串
            field (str): 字段名称
            
        Returns:
            str: 格式化的日期时间字符串
        """
        try:
            # 去除多余的下划线和空格
            value = re.sub(r'[_\s]+', ' ', value).strip()
            
            # 尝试解析不同格式的日期
            # 格式：2023年10月24日
            match = re.search(r'(\d{4})[年-](\d{1,2})[月-](\d{1,2})日?', value)
            if match:
                year = match.group(1)
                month = match.group(2).zfill(2)
                day = match.group(3).zfill(2)
                return f"{year}.{month}.{day}"
            
            # 格式：2023-10-24
            match = re.search(r'(\d{4})-(\d{1,2})-(\d{1,2})', value)
            if match:
                year = match.group(1)
                month = match.group(2).zfill(2)
                day = match.group(3).zfill(2)
                return f"{year}.{month}.{day}"
            
            # 格式：2023.10.24
            match = re.search(r'(\d{4})\.(\d{1,2})\.(\d{1,2})', value)
            if match:
                year = match.group(1)
                month = match.group(2).zfill(2)
                day = match.group(3).zfill(2)
                return f"{year}.{month}.{day}"
            
            # 格式：2023/10/24
            match = re.search(r'(\d{4})/(\d{1,2})/(\d{1,2})', value)
            if match:
                year = match.group(1)
                month = match.group(2).zfill(2)
                day = match.group(3).zfill(2)
                return f"{year}.{month}.{day}"
            
            # 如果包含时间
            match = re.search(r'(\d{4})[年-](\d{1,2})[月-](\d{1,2})日?\s+(\d{1,2}):(\d{1,2})', value)
            if match:
                year = match.group(1)
                month = match.group(2).zfill(2)
                day = match.group(3).zfill(2)
                hour = match.group(4).zfill(2)
                minute = match.group(5).zfill(2)
                return f"{year}.{month}.{day} {hour}:{minute}"
            
            # 如果无法解析，返回原始值
            return value
        except Exception as e:
            logger.error(f"解析日期时间出错: {str(e)}")
            return value
    
    def clean_data(self, data):
        """
        清理提取的数据，去除多余的下划线和空格
        
        Args:
            data (dict): 提取的数据
            
        Returns:
            dict: 清理后的数据
        """
        cleaned_data = {}
        for field, value in data.items():
            if isinstance(value, str):
                # 去除多余的下划线和空格
                cleaned_value = re.sub(r'[_\s]+', ' ', value).strip()
                cleaned_data[field] = cleaned_value
            else:
                cleaned_data[field] = value
        
        return cleaned_data
    
    def convert_to_dataframe(self, data_dict):
        """
        将提取的数据转换为DataFrame，并按指定顺序排列字段
        
        Args:
            data_dict (dict): 提取的数据字典
            
        Returns:
            pandas.DataFrame: 包含数据的DataFrame
        """
        # 创建只有一行的DataFrame
        df = pd.DataFrame([data_dict])
        
        # 定义字段顺序
        field_order = [
            '姓名', '性别', '年龄', '出生日期', '国籍', '出生地址', '民族', 
            '证件类别', '证件号码', '职业', '入院途径', '入院时间', '入院科别', 
            '入院病房', '转科科别', '出院时间', '出院科别', '出院病房', 
            '实际住院', '门诊诊断名称', '门诊诊断编码', '入院诊断名称', 
            '入院诊断编码', '主要诊断', '主要诊断编码', '其他诊断', 
            '其他诊断编码', '损伤中毒的外部原因', '病理诊断名称', 
            '病理诊断编码', '病理描述', '病理号', '是否有药物过敏', 
            '过敏药物', '是否进行了死亡患者尸检'
        ]
        
        # 重新排序列（如果存在）
        existing_columns = [col for col in field_order if col in df.columns]
        other_columns = [col for col in df.columns if col not in field_order]
        df = df[existing_columns + other_columns]
        
        return df
    
    def extract_and_convert(self, docx_path):
        """
        从Word文档提取数据并转换为DataFrame
        
        Args:
            docx_path (str): Word文档的路径
            
        Returns:
            pandas.DataFrame: 包含提取数据的DataFrame
        """
        data = self.extract_data_from_docx(docx_path)
        return self.convert_to_dataframe(data)

    def extract_data_from_docx(self, docx_path):
        """
        从Word文档中提取数据
        
        Args:
            docx_path (str): Word文档的路径
            
        Returns:
            dict: 提取的数据字段及其值
        """
        try:
            # 打开Word文档
            doc = Document(docx_path)
            
            # 获取文档全文
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # 表格数据也很重要
            table_data = {}
            for table in doc.tables:
                # 处理表格，提取表头和数据
                headers = []
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip())
                
                # 处理表格数据行
                for row in table.rows[1:]:  # 跳过表头行
                    row_data = {}
                    for i, cell in enumerate(row.cells):
                        if i < len(headers):
                            header = headers[i]
                            value = cell.text.strip()
                            row_data[header] = value
                    
                    # 将行数据添加到表格数据中
                    for header, value in row_data.items():
                        if header and value:  # 只处理非空的表头和值
                            # 尝试将表头映射到标准字段名
                            standardized_header = self._map_header_to_field(header)
                            if standardized_header:
                                table_data[standardized_header] = value
                
                # 尝试直接从表格结构中提取特定字段
                self._extract_from_table_structure(table, table_data)
                
                # 收集表格行文本用于常规提取
                for row in table.rows:
                    row_text = ' '.join([cell.text for cell in row.cells])
                    full_text.append(row_text)
            
            # 将所有文本连接为一个字符串
            text = '\n'.join(full_text)
            
            # 使用extract_data方法处理提取的文本
            data = self.extract_data(text)
            
            # 合并表格专门提取的数据
            for field, value in table_data.items():
                if field not in data or not data[field]:
                    data[field] = value
            
            return data
            
        except Exception as e:
            logger.error(f"提取数据时出错: {str(e)}")
            return {}
            
    def _map_header_to_field(self, header):
        """
        将表格表头映射到标准字段名
        
        Args:
            header (str): 表格表头
            
        Returns:
            str: 标准字段名，如果没有匹配则返回None
        """
        # 表头映射字典
        header_mapping = {
            '入院诊断名称': '入院诊断名称',
            '入院诊断': '入院诊断名称',
            '入院诊断编码': '入院诊断编码',
            '主要诊断': '主要诊断',
            '其他诊断': '其他诊断',
            '损伤中毒的外部原因': '损伤中毒的外部原因',
            '损伤、中毒的外部原因': '损伤中毒的外部原因',
            '病理诊断名称': '病理诊断名称',
            '病理诊断': '病理诊断名称',
            '病理诊断编码': '病理诊断编码',
            '病理描述': '病理描述',
            '病理号': '病理号',
            '过敏药物': '过敏药物',
            '过敏药物源': '过敏药物'
        }
        
        # 尝试直接匹配
        if header in header_mapping:
            return header_mapping[header]
        
        # 尝试模糊匹配
        for key, value in header_mapping.items():
            if key in header:
                return value
        
        return None

    def _extract_from_table_structure(self, table, data):
        """
        直接从表格结构中提取数据，针对特定表格布局
        
        Args:
            table: Word表格对象
            data (dict): 要更新的数据字典
        """
        try:
            # 检查表格是否有足够的行和列
            if len(table.rows) < 2 or len(table.rows[0].cells) < 2:
                return
                
            # 获取表头
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            # 特定字段在表格中的列索引映射
            field_column_mapping = {
                '出生地址': None,
                '入院途径': None,
                '入院诊断名称': None,
                '入院诊断编码': None,
                '主要诊断': None,
                '其他诊断': None,
                '损伤中毒的外部原因': None,
                '病理诊断名称': None,
                '病理诊断编码': None,
                '病理描述': None,
                '病理号': None
            }
            
            # 查找字段对应的列索引
            for i, header in enumerate(headers):
                header_lower = header.lower()
                if '出生地' in header_lower:
                    field_column_mapping['出生地址'] = i
                elif '入院途径' in header_lower:
                    field_column_mapping['入院途径'] = i
                elif '入院诊断' in header_lower and ('名称' in header_lower or '诊断名' in header_lower):
                    field_column_mapping['入院诊断名称'] = i
                elif '入院诊断' in header_lower and '编码' in header_lower:
                    field_column_mapping['入院诊断编码'] = i
                elif '主要诊断' in header_lower:
                    field_column_mapping['主要诊断'] = i
                elif '其他诊断' in header_lower:
                    field_column_mapping['其他诊断'] = i
                elif ('损伤' in header_lower and '外部原因' in header_lower) or ('中毒' in header_lower and '外部原因' in header_lower):
                    field_column_mapping['损伤中毒的外部原因'] = i
                elif '病理诊断' in header_lower and ('名称' in header_lower or '诊断名' in header_lower):
                    field_column_mapping['病理诊断名称'] = i
                elif '病理诊断' in header_lower and '编码' in header_lower:
                    field_column_mapping['病理诊断编码'] = i
                elif '病理描述' in header_lower:
                    field_column_mapping['病理描述'] = i
                elif '病理号' in header_lower:
                    field_column_mapping['病理号'] = i
            
            # 从数据行中提取字段值
            for row in table.rows[1:]:  # 跳过表头行
                for field, col_idx in field_column_mapping.items():
                    if col_idx is not None and col_idx < len(row.cells):
                        value = row.cells[col_idx].text.strip()
                        if value:
                            # 处理特殊情况
                            if field == '出生地址':
                                # 处理出生地址中可能包含的民族和籍贯信息
                                ethnicity_match = re.search(r'民族[:：\s]*([\u4e00-\u9fa5]+族?)', value)
                                if ethnicity_match:
                                    ethnicity = ethnicity_match.group(1)
                                    value = value.replace(ethnicity_match.group(0), '').strip()
                                    if '民族' not in data or not data['民族']:
                                        data['民族'] = ethnicity
                                
                                origin_match = re.search(r'籍贯[:：\s]*([\u4e00-\u9fa5]+)', value)
                                if origin_match:
                                    value = value.replace(origin_match.group(0), '').strip()
                                
                                data[field] = value
                                
                            elif field == '入院途径':
                                # 处理可能包含说明的入院途径，如"2 1.急诊，2门诊，3其他医疗机构转入，9其他"
                                # 先尝试提取开头的数字作为代码
                                code_match = re.match(r'^(\d+)', value)
                                if code_match:
                                    code = code_match.group(1)
                                    # 使用代码映射
                                    if code in self.admission_path_mapping:
                                        data[field] = self.admission_path_mapping[code]
                                    else:
                                        data[field] = value
                                else:
                                    # 如果没有找到代码，尝试使用整个值进行映射
                                    data[field] = self.admission_path_mapping.get(value, value)
                                
                            elif field == '入院诊断名称':
                                # 检查是否包含编码
                                code_match = re.search(r'([A-Z]\d+\.\d+|[A-Z]\d+)', value)
                                if code_match:
                                    code = code_match.group(0)
                                    # 从诊断名称中移除编码
                                    clean_value = value.replace(code, '').strip()
                                    data[field] = clean_value
                                    
                                    # 如果没有编码字段，添加它
                                    if '入院诊断编码' not in data or not data['入院诊断编码']:
                                        data['入院诊断编码'] = code
                                else:
                                    data[field] = value
                            elif field == '病理诊断名称':
                                # 检查是否包含编码
                                code_match = re.search(r'([A-Z]\d+\.\d+|[A-Z]\d+)', value)
                                if code_match:
                                    code = code_match.group(0)
                                    # 从诊断名称中移除编码
                                    clean_value = value.replace(code, '').strip()
                                    data[field] = clean_value
                                    
                                    # 如果没有编码字段，添加它
                                    if '病理诊断编码' not in data or not data['病理诊断编码']:
                                        data['病理诊断编码'] = code
                                else:
                                    data[field] = value
                            elif field == '病理描述':
                                # 检查是否包含病理号
                                path_num_match = re.search(r'病理号[：:]\s*([A-Z0-9-]+)', value)
                                if path_num_match:
                                    path_num = path_num_match.group(0)
                                    # 从描述中移除病理号
                                    clean_value = value.replace(path_num, '').strip()
                                    data[field] = clean_value
                                    
                                    # 如果没有病理号字段，添加它
                                    if '病理号' not in data or not data['病理号']:
                                        data['病理号'] = path_num_match.group(1)
                                else:
                                    data[field] = value
                            else:
                                data[field] = value
        except Exception as e:
            logger.error(f"从表格结构提取数据时出错: {str(e)}")
            # 错误不应该中断整个处理流程

    def _clean_merged_fields(self, data):
        """
        清理可能错误合并的字段
        
        Args:
            data (dict): 提取的数据字典
        """
        # 检查所有字段的值，查找可能包含下一个字段名称的值
        fields_to_check = [
            ('姓名', '性别'),
            ('性别', '年龄'),
            ('年龄', '出生日期'),
            ('出生日期', '国籍'),
            ('国籍', '出生地址'),
            ('出生地址', '民族'),
            ('民族', '证件类别'),
            ('证件类别', '证件号码'),
            ('证件号码', '职业'),
            ('职业', '入院途径'),
            ('入院途径', '入院时间'),
            ('入院时间', '入院科别'),
            ('入院科别', '入院病房'),
            ('入院病房', '转科科别'),
            ('转科科别', '出院时间'),
            ('出院时间', '出院科别'),
            ('出院科别', '出院病房'),
            ('出院病房', '实际住院'),
            ('实际住院', '入院诊断名称'),
            ('入院诊断名称', '入院诊断编码'),
            ('入院诊断编码', '主要诊断'),
            ('主要诊断', '主要诊断编码'),
            ('主要诊断编码', '其他诊断'),
            ('其他诊断', '其他诊断编码'),
            ('其他诊断编码', '损伤中毒的外部原因'),
            ('损伤中毒的外部原因', '病理诊断名称'),
            ('病理诊断名称', '病理诊断编码'),
            ('病理诊断编码', '病理描述'),
            ('病理描述', '病理号'),
            ('病理号', '是否有药物过敏'),
            ('是否有药物过敏', '过敏药物'),
            ('过敏药物', '是否进行了死亡患者尸检')
        ]
        
        for current_field, next_field in fields_to_check:
            # 如果当前字段存在且下一个字段在字段列表中
            if current_field in data and next_field in self.field_name_mapping.values():
                value = data[current_field]
                # 检查值是否以下一个字段名开头
                for next_field_alias in [k for k, v in self.field_name_mapping.items() if v == next_field]:
                    if value and isinstance(value, str):
                        # 如果当前字段的值包含下一个字段的名称，则可能是错误合并的
                        if value.startswith(next_field_alias) or f"；{next_field_alias}" in value or f";{next_field_alias}" in value:
                            # 将当前字段的值设置为空字符串
                            data[current_field] = ""
                            break

    def _fix_known_errors(self, data):
        """
        修复已知的错误模式
        
        Args:
            data (dict): 提取的数据字典
        """
        # 修复"姓名"字段值为"；性别2"的情况
        if '姓名' in data and data['姓名'] and isinstance(data['姓名'], str):
            if '性别' in data['姓名'] or '；性别' in data['姓名'] or ';性别' in data['姓名']:
                data['姓名'] = ""
        
        # 修复"证件号码"字段值为"；职业"的情况
        if '证件号码' in data and data['证件号码'] and isinstance(data['证件号码'], str):
            if '职业' in data['证件号码'] or '；职业' in data['证件号码'] or ';职业' in data['证件号码']:
                data['证件号码'] = ""
                
        # 检查所有字段，查找可能包含其他字段名称的值
        fields_to_check = [
            ('姓名', ['性别']),
            ('性别', ['年龄']),
            ('年龄', ['出生日期']),
            ('出生日期', ['国籍']),
            ('国籍', ['出生地址', '出生地']),
            ('出生地址', ['民族']),
            ('民族', ['证件类别']),
            ('证件类别', ['证件号码', '身份证号']),
            ('证件号码', ['职业']),
            ('职业', ['入院途径']),
            ('入院途径', ['入院时间', '入院日期']),
            ('入院时间', ['入院科别', '科室']),
            ('入院科别', ['入院病房', '病房', '床号']),
            ('入院病房', ['转科科别']),
            ('转科科别', ['出院时间', '出院日期']),
            ('出院时间', ['出院科别']),
            ('出院科别', ['出院病房']),
            ('出院病房', ['实际住院', '住院天数']),
            ('实际住院', ['入院诊断名称', '入院诊断']),
            ('入院诊断名称', ['入院诊断编码']),
            ('入院诊断编码', ['主要诊断']),
            ('主要诊断', ['主要诊断编码']),
            ('主要诊断编码', ['其他诊断']),
            ('其他诊断', ['其他诊断编码']),
            ('其他诊断编码', ['损伤中毒的外部原因', '损伤、中毒的外部原因']),
            ('损伤中毒的外部原因', ['病理诊断名称', '病理诊断']),
            ('病理诊断名称', ['病理诊断编码']),
            ('病理诊断编码', ['病理描述']),
            ('病理描述', ['病理号']),
            ('病理号', ['是否有药物过敏']),
            ('是否有药物过敏', ['过敏药物', '过敏药物源']),
            ('过敏药物', ['是否进行了死亡患者尸检'])
        ]
        
        for current_field, next_fields in fields_to_check:
            if current_field in data and data[current_field] and isinstance(data[current_field], str):
                value = data[current_field]
                for next_field in next_fields:
                    # 检查当前字段的值是否包含下一个字段的名称
                    if next_field in value or f'；{next_field}' in value or f';{next_field}' in value:
                        # 如果包含，则将当前字段的值设置为空
                        data[current_field] = ""
                        break
                        
        # 检查并修复可能的格式问题
        for field in data:
            if isinstance(data[field], str):
                # 移除值中可能的分号前缀
                if data[field].startswith('；') or data[field].startswith(';'):
                    data[field] = data[field][1:].strip()
                    
                # 如果值只包含下划线或空格，则设为空字符串
                if re.match(r'^[\s_]+$', data[field]):
                    data[field] = ""


class DocxResultGenerator:
    """
    生成带有错误标记的Word文档
    """
    def __init__(self):
        """初始化结果生成器"""
        pass
    
    def highlight_errors(self, docx_path, errors, output_path=None):
        """
        根据错误信息在Word文档中高亮显示错误内容
        
        Args:
            docx_path (str): 原始Word文档的路径
            errors (list): 错误信息列表
            output_path (str, optional): 输出文档的路径，默认为None（自动生成）
            
        Returns:
            str: 输出文档的路径
        """
        try:
            # 打开原始文档
            doc = Document(docx_path)
            
            # 如果未指定输出路径，则自动生成
            if output_path is None:
                file_dir = os.path.dirname(docx_path)
                file_name = os.path.basename(docx_path)
                name, ext = os.path.splitext(file_name)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = os.path.join(file_dir, f"{name}_检查结果_{timestamp}{ext}")
            
            # 获取错误字段列表
            error_fields = set()
            for error in errors:
                if 'rule_name' in error and 'message' in error:
                    # 从错误消息中提取字段名
                    rule_name = error['rule_name']
                    if '缺' in rule_name or '漏' in rule_name:
                        # 缺项规则，直接从规则名提取字段
                        field_match = re.search(r'([^缺漏]+)[缺漏]', rule_name)
                        if field_match:
                            error_fields.add(field_match.group(1))
                    
                    # 从消息中提取字段名
                    for field in self.field_patterns.keys():
                        if field in error['message']:
                            error_fields.add(field)
            
            # 在段落中标记错误
            for para in doc.paragraphs:
                self._mark_errors_in_text(para, error_fields)
            
            # 在表格中标记错误
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._mark_errors_in_text(para, error_fields)
            
            # 保存结果文档
            doc.save(output_path)
            return output_path
        
        except Exception as e:
            logger.error(f"高亮错误时出错: {str(e)}")
            return None
    
    def _mark_errors_in_text(self, paragraph, error_fields):
        """
        在段落文本中标记错误
        
        Args:
            paragraph: 文档段落对象
            error_fields (set): 错误字段集合
        """
        text = paragraph.text
        if not text:
            return
        
        # 检查段落是否包含错误字段
        for field in error_fields:
            # 更精确的字段匹配模式
            field_pattern = f"{field}[：:_\\s]*"
            match = re.search(field_pattern, text)
            if match:
                # 获取字段在文本中的位置
                start_pos = match.start()
                end_pos = match.end()
                
                # 清除段落中的所有运行
                for _ in range(len(paragraph.runs)):
                    paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
                
                # 分三部分添加文本：字段前、字段、字段后
                if start_pos > 0:
                    paragraph.add_run(text[:start_pos])
                
                # 添加红色高亮的字段部分
                field_run = paragraph.add_run(text[start_pos:end_pos])
                from docx.shared import RGBColor
                field_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
                
                # 查找字段对应的值
                value_match = None
                # 尝试查找字段后面的值
                if end_pos < len(text):
                    # 查找到下一个字段或行尾
                    next_field_match = None
                    for next_field in error_fields:
                        if next_field != field:
                            next_match = re.search(f"{next_field}[：:_\\s]*", text[end_pos:])
                            if next_match and (next_field_match is None or next_match.start() < next_field_match.start()):
                                next_field_match = next_match
                    
                    if next_field_match:
                        value_text = text[end_pos:end_pos+next_field_match.start()]
                    else:
                        value_text = text[end_pos:]
                    
                    # 添加红色高亮的值部分
                    value_run = paragraph.add_run(value_text)
                    value_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
                    
                    # 如果有下一个字段，添加剩余文本
                    if next_field_match:
                        paragraph.add_run(text[end_pos+next_field_match.start():])
                
                # 已处理此段落，退出循环
                return
    
    @property
    def field_patterns(self):
        """获取字段模式"""
        return {
            '姓名': r'姓名[_\s]*([^\s_]+)[_\s]*',
            '性别': r'性别[_\s]*([男女12])[_\s]*',
            '年龄': r'年龄[_\s]*(\d+)[_\s]*',
            '出生日期': r'出生日期[_\s]*(\d{4})[年_\s]+(\d{1,2})[月_\s]+(\d{1,2})[日_\s]*',
            '国籍': r'国籍[_\s]*([^\s_]+)[_\s]*',
            '出生地址': r'出生地址[_\s]*([^\n]+)[_\s]*',
            '民族': r'民族[_\s]*([^\s_]+)[_\s]*',
            '证件类别': r'证件类别[_\s]*([^\s_]+)[_\s]*',
            '证件号码': r'(?:证件号码|身份证号)[_\s]*([0-9Xx]{15,18}|[^\s_]+)[_\s]*',
            '入院途径': r'入院途径[_\s]*([^\s_]+)[_\s]*',
            '入院时间': r'(?:入院时间|入院日期)[_\s]*(\d{4})[-年_\s]+(\d{1,2})[-月_\s]+(\d{1,2})(?:[-日_\s]+(\d{1,2}):(\d{1,2}))?[_\s]*',
            '入院科别': r'(?:入院科别|科室)[_\s]*([^\s_]+)[_\s]*',
            '入院病房': r'(?:入院病房|病房|床号)[_\s]*([^\s_]+)[_\s]*',
            '转科科别': r'转科科别[_\s]*([^\s_]+)[_\s]*',
            '出院时间': r'(?:出院时间|出院日期)[_\s]*(\d{4})[-年_\s]+(\d{1,2})[-月_\s]+(\d{1,2})(?:[-日_\s]+(\d{1,2}):(\d{1,2}))?[_\s]*',
            '出院科别': r'出院科别[_\s]*([^\s_]+)[_\s]*',
            '出院病房': r'出院病房[_\s]*([^\s_]+)[_\s]*',
            '实际住院': r'(?:实际住院|住院天数)[_\s]*(\d+)[天_\s]*',
            '门诊诊断名称': r'(?:门诊诊断名称|门急诊诊断名称)[_\s]*([^\n]+)[_\s]*',
            '门诊诊断编码': r'(?:门诊诊断编码|门急诊诊断编码)[_\s]*([^\n]+)[_\s]*',
            '入院诊断名称': r'入院诊断(?:名称)?[_\s]*([^\n]+)[_\s]*',
            '入院诊断编码': r'入院诊断编码[_\s]*([^\n]+)[_\s]*',
            '主要诊断': r'主要诊断[_\s]*([^\n]+)[_\s]*',
            '其他诊断': r'其他诊断[_\s]*([^\n]+)[_\s]*',
            '损伤中毒的外部原因': r'(?:损伤中毒的外部原因|损伤、中毒的外部原因)[_\s]*([^\n]+)[_\s]*',
            '病理诊断名称': r'病理诊断(?:名称)?[_\s]*([^\n]+)[_\s]*',
            '病理诊断编码': r'病理诊断编码[_\s]*([^\n]+)[_\s]*',
            '病理描述': r'病理描述[_\s]*([^\n]+)[_\s]*',
            '病理号': r'病理号[_\s]*([^\s_]+)[_\s]*',
            '是否有药物过敏': r'是否有药物过敏[_\s]*([是否有无YyNn01是的不是]+)[_\s]*',
            '过敏药物': r'(?:过敏药物|过敏药物源)[_\s]*([^\n]*)[_\s]*',
            '是否进行了死亡患者尸检': r'是否进行了死亡患者尸检[_\s]*([是否有无YyNn01是的不是]+)[_\s]*',
        } 