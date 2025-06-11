from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_from_directory, send_file
import os
import json
import pandas as pd
import numpy as np
from datetime import datetime
import re
from werkzeug.utils import secure_filename
from medical_entities import get_medical_entities, save_medical_entities, recognize_entities
from export_medical_records import export_medical_records, export_patients, export_admissions
from text_to_excel import parse_medical_text
# 导入大模型命名实体识别模块
from llm_ner import get_llm_config, save_llm_config, recognize_entities_with_api, calculate_entity_statistics, recognize_entities_with_rules
from docx_data_check import DocxDataExtractor, DocxResultGenerator
import html
import logging

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 初始化Flask应用
app = Flask(__name__)
app.secret_key = 'your_secret_key'  # 用于会话安全的密钥，生产环境应使用强随机密钥

# 创建必要的存储目录
os.makedirs('data', exist_ok=True)    # 用于存储规则和映射数据
os.makedirs('uploads', exist_ok=True)  # 用于存储上传的文件
os.makedirs('excel_data', exist_ok=True)  # 用于存储生成的Excel文件

# 文件路径常量定义
RULES_FILE = 'data/rules.json'  # 规则文件路径
DIAGNOSIS_DEPT_MAPPING_FILE = 'data/diagnosis_department_mapping.json'  # 科室与诊断映射文件路径
MEDICAL_ENTITIES_FILE = 'data/medical_entities.json'  # 医学实体字典文件路径
DB_CONFIG_FILE = 'data/db_config.json'  # 数据库配置文件路径
LLM_CONFIG_FILE = 'data/llm_config.json'  # LLM配置文件路径

# 初始化规则文件（如果不存在）
if not os.path.exists(RULES_FILE):
    with open(RULES_FILE, 'w', encoding='utf-8') as f:
        json.dump([], f, ensure_ascii=False)

# 初始化医学实体字典文件（如果不存在）
if not os.path.exists(MEDICAL_ENTITIES_FILE):
    medical_entities = {
        "疾病": ["肺癌", "支气管哮喘", "冠心病", "高血压", "糖尿病", "肝炎", "肺炎"],
        "症状": ["发热", "咳嗽", "胸痛", "头痛", "腹痛", "呕吐", "腹泻"],
        "检查": ["血常规", "尿常规", "肝功能", "CT检查", "核磁共振", "X光检查", "超声检查"],
        "治疗": ["手术", "药物治疗", "放疗", "化疗", "物理治疗", "心理治疗", "康复治疗"],
        "药物": ["青霉素", "阿莫西林", "头孢", "阿司匹林", "布洛芬", "泼尼松", "胰岛素"]
    }
    with open(MEDICAL_ENTITIES_FILE, 'w', encoding='utf-8') as f:
        json.dump(medical_entities, ensure_ascii=False, indent=2, fp=f)

# 初始化数据库配置文件（如果不存在）
if not os.path.exists(DB_CONFIG_FILE):
    db_config = {
        'host': 'localhost',
        'user': 'root',
        'password': '',
        'database': 'hospital_emr',
        'charset': 'utf8mb4'
    }
    with open(DB_CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(db_config, ensure_ascii=False, indent=2, fp=f)

# 初始化大模型配置文件（如果不存在）
if not os.path.exists(LLM_CONFIG_FILE):
    llm_config = {
        'api_mode': False,
        'api_type': 'deepseek',
        'api_key': '',
        'api_url': ''
    }
    with open(LLM_CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(llm_config, ensure_ascii=False, indent=2, fp=f)

# 获取所有规则
def get_rules():
    """
    从规则文件中读取所有质控规则
    
    读取JSON格式的规则文件，解析其中定义的质控规则列表。
    如果文件不存在或格式错误，则返回空列表并记录错误。
    
    Returns:
        list: 质控规则列表，每个规则为一个字典，包含规则ID、名称、类型、条件等信息
    """
    try:
        with open(RULES_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"读取规则文件出错: {str(e)}")
        return []

# 保存规则
def save_rules(rules):
    """
    将规则列表保存到规则文件
    
    将质控规则列表序列化为JSON格式并保存到规则文件中。
    如果保存过程中出现错误，则记录错误信息。
    
    Args:
        rules (list): 要保存的规则列表，每个规则为一个字典
    """
    try:
        with open(RULES_FILE, 'w', encoding='utf-8') as f:
            json.dump(rules, ensure_ascii=False, indent=2, fp=f)
    except Exception as e:
        print(f"保存规则文件出错: {str(e)}")

# 获取科室与诊断映射
def get_diagnosis_dept_mapping():
    """
    从映射文件中读取科室与诊断的对应关系
    
    Returns:
        dict: 科室与诊断的映射字典，如果文件不存在或格式错误则返回空字典
    """
    try:
        with open(DIAGNOSIS_DEPT_MAPPING_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"读取科室诊断映射文件出错: {str(e)}")
        return {}

# 获取医学实体字典
def get_medical_entities():
    """
    从医学实体字典文件中读取医学实体数据
    
    Returns:
        dict: 医学实体字典，如果文件不存在或格式错误则返回空字典
    """
    try:
        with open(MEDICAL_ENTITIES_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"读取医学实体字典文件出错: {str(e)}")
        return {}

# 保存医学实体字典
def save_medical_entities(entities):
    """
    将医学实体字典保存到文件
    
    Args:
        entities (dict): 要保存的医学实体字典
    """
    try:
        with open(MEDICAL_ENTITIES_FILE, 'w', encoding='utf-8') as f:
            json.dump(entities, ensure_ascii=False, indent=2, fp=f)
    except Exception as e:
        print(f"保存医学实体字典文件出错: {str(e)}")

# 获取数据库配置
def get_db_config():
    """
    从配置文件中读取数据库配置
    
    Returns:
        dict: 数据库配置字典，如果文件不存在或格式错误则返回默认配置
    """
    try:
        with open(DB_CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"读取数据库配置文件出错: {str(e)}")
        return {
            'host': 'localhost',
            'user': 'root',
            'password': '',
            'database': 'hospital_emr',
            'charset': 'utf8mb4'
        }

# 保存数据库配置
def save_db_config(config):
    """
    将数据库配置保存到文件
    
    Args:
        config (dict): 要保存的数据库配置
    """
    try:
        with open(DB_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, ensure_ascii=False, indent=2, fp=f)
    except Exception as e:
        print(f"保存数据库配置文件出错: {str(e)}")

# 路由定义部分
# 首页
@app.route('/')
def index():
    """首页路由，显示系统主页"""
    return render_template('index.html')

# 规则管理页面
@app.route('/rules')
def rules_page():
    """规则管理页面路由，显示规则管理界面"""
    rules = get_rules()
    return render_template('rules.html', rules=rules)

# 数据库转Excel页面
@app.route('/database_to_excel')
def database_to_excel():
    """数据库转Excel页面路由，显示数据库导出界面"""
    return render_template('database_to_excel.html')

# 处理数据库导出请求
@app.route('/export_database', methods=['POST'])
def export_database():
    """处理数据库导出请求的路由"""
    try:
        export_type = request.form.get('export_type')
        db_type = request.form.get('db_type')
        host = request.form.get('host')
        port = request.form.get('port')
        username = request.form.get('username')
        password = request.form.get('password')
        database = request.form.get('database')
        
        if not all([export_type, db_type, host, username, database]):
            flash('请填写所有必填字段')
            return redirect(url_for('database_to_excel'))
        
        # 创建数据库连接
        try:
            # 处理端口
            if port:
                port = int(port)
            else:
                # 根据数据库类型设置默认端口
                if db_type == 'mysql':
                    port = 3306
                elif db_type == 'postgresql':
                    port = 5432
                elif db_type == 'sqlserver':
                    port = 1433
                elif db_type == 'oracle':
                    port = 1521
                # SQLite不需要端口
            
            # 根据数据库类型创建连接
            if db_type == 'mysql':
                import mysql.connector
                conn = mysql.connector.connect(
                    host=host,
                    port=port,
                    user=username,
                    password=password,
                    database=database
                )
            elif db_type == 'postgresql':
                import psycopg2
                conn = psycopg2.connect(
                    host=host,
                    port=port,
                    user=username,
                    password=password,
                    dbname=database
                )
            elif db_type == 'sqlserver':
                import pyodbc
                conn = pyodbc.connect(
                    f'DRIVER={{SQL Server}};SERVER={host},{port};DATABASE={database};UID={username};PWD={password}'
                )
            elif db_type == 'oracle':
                import cx_Oracle
                conn = cx_Oracle.connect(
                    f'{username}/{password}@{host}:{port}/{database}'
                )
            elif db_type == 'sqlite':
                import sqlite3
                conn = sqlite3.connect(database)
            else:
                flash('不支持的数据库类型')
                return redirect(url_for('database_to_excel'))
            
            # 根据导出类型调用相应的函数
            if export_type == 'patients':
                result_file = export_patients(conn)
                flash('患者信息导出成功')
            elif export_type == 'admissions':
                result_file = export_admissions(conn)
                flash('住院记录导出成功')
            elif export_type == 'medical_records':
                result_file = export_medical_records(conn)
                flash('病案首页导出成功')
            else:
                flash('不支持的导出类型')
                return redirect(url_for('database_to_excel'))
        
        except Exception as e:
            flash(f'数据库连接或导出错误: {str(e)}')
            return redirect(url_for('database_to_excel'))
        
        # 关闭连接
        conn.close()
        
        # 返回下载链接
        return send_file(result_file, as_attachment=True)
    except Exception as e:
        flash(f'导出数据库出错: {str(e)}')
        return redirect(url_for('database_to_excel'))

# 文本转Excel页面
@app.route('/text_to_excel')
def text_to_excel_page():
    """文本转Excel页面路由，显示文本转换界面"""
    return render_template('text_to_excel.html')

# 处理文本转Excel请求
@app.route('/process_text_to_excel', methods=['POST'])
def process_text_to_excel():
    """
    处理医疗文本转Excel请求的路由
    
    接收用户提交的医疗记录文本，解析其中的结构化数据，
    并生成标准化的Excel文件供用户下载。支持文本输入和文件上传两种方式。
    
    Returns:
        str: 渲染后的文本转Excel结果页面HTML，包含下载链接
    """
    try:
        # 获取表单数据
        text = request.form.get('text', '')
        
        if not text:
            flash('请输入文本')
            return redirect(url_for('text_to_excel_page'))
        
        # 解析文本
        data = parse_medical_text(text)
        
        if not data:
            flash('无法解析文本，请检查文本格式')
            return redirect(url_for('text_to_excel_page'))
        
        # 转换为DataFrame
        df = pd.DataFrame(data)
        
        # 生成Excel文件
        if not os.path.exists('excel_data'):
            os.makedirs('excel_data')
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"医疗记录_{timestamp}.xlsx"
        filepath = os.path.join('excel_data', filename)
        
        df.to_excel(filepath, index=False, engine='openpyxl')
        
        # 生成HTML表格预览
        html_table = df.to_html(classes='table table-striped table-bordered', index=False)
        
        # 返回预览页面
        return render_template('text_to_excel_result.html', 
                               filename=filename, 
                               data=html_table, 
                               record_count=len(data))
    except Exception as e:
        flash(f'处理文本转Excel出错: {str(e)}')
        return redirect(url_for('text_to_excel_page'))

# 实体识别页面
@app.route('/entity_recognition')
def entity_recognition_page():
    """实体识别页面路由，显示实体识别界面"""
    entities = get_medical_entities()
    return render_template('entity_recognition.html', entities=entities)

# 处理医学文本实体识别请求
@app.route('/recognize_entities', methods=['POST'])
def recognize_entities():
    """
    处理医学实体识别请求的路由
    
    接收用户提交的医学文本，使用词典匹配方法识别其中的医学实体，
    并返回高亮显示识别结果的页面。支持文本输入和文件上传两种方式。
    
    Returns:
        str: 渲染后的实体识别结果页面HTML
    """
    try:
        # 获取表单数据
        text = request.form.get('text', '')
        
        if not text:
            flash('请输入医学文本')
            return redirect(url_for('entity_recognition_page'))
        
        # 获取医学实体字典
        entities_dict = get_medical_entities()
        
        # 识别实体
        recognized_entities = {}
        entity_statistics = {}  # 添加实体统计字典
        
        for entity_type, entity_list in entities_dict.items():
            found_entities = []
            entity_counts = {}  # 每个实体类型下各实体的计数
            
            for entity in entity_list:
                if entity in text:
                    # 找出所有匹配位置
                    positions = [m.start() for m in re.finditer(re.escape(entity), text)]
                    count = len(positions)  # 该实体出现的次数
                    
                    if count > 0:
                        entity_counts[entity] = count  # 记录该实体的出现次数
                        
                    for pos in positions:
                        found_entities.append({
                            'entity': entity,
                            'position': pos,
                            'context': text[max(0, pos-10):min(len(text), pos+len(entity)+10)]
                        })
            
            if found_entities:
                recognized_entities[entity_type] = found_entities
                entity_statistics[entity_type] = entity_counts  # 添加该类型的实体统计
        
        # 准备高亮显示的文本
        highlighted_text = text
        for entity_type, entities in recognized_entities.items():
            for entity_info in sorted(entities, key=lambda x: x['position'], reverse=True):
                entity = entity_info['entity']
                pos = entity_info['position']
                highlighted_text = (
                    highlighted_text[:pos] + 
                    f'<span class="entity-highlight {entity_type.lower()}" title="{entity_type}">{entity}</span>' + 
                    highlighted_text[pos+len(entity):]
                )
        
        return render_template('entity_recognition_result.html', 
                               original_text=text,
                               highlighted_text=highlighted_text,
                               recognized_entities=recognized_entities,
                               entity_statistics=entity_statistics)  # 传递实体统计数据到模板
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"实体识别错误: {str(e)}\n{error_details}")
        flash(f'实体识别错误: {str(e)}')
        return redirect(url_for('entity_recognition_page'))

# 添加医学实体
@app.route('/add_entity', methods=['POST'])
def add_entity():
    """
    添加新医学实体的路由
    """
    try:
        entity_type = request.form.get('entity_type')
        entity_name = request.form.get('entity_name')
        
        if not entity_type or not entity_name:
            flash('实体类型和名称不能为空')
            return redirect(url_for('entity_recognition_page'))
        
        # 获取现有实体
        entities = get_medical_entities()
        
        # 如果实体类型不存在，创建新类型
        if entity_type not in entities:
            entities[entity_type] = []
        
        # 如果实体名称不存在，添加到列表
        if entity_name not in entities[entity_type]:
            entities[entity_type].append(entity_name)
            save_medical_entities(entities)
            flash(f'成功添加实体: {entity_name} (类型: {entity_type})')
        else:
            flash(f'实体已存在: {entity_name} (类型: {entity_type})')
        
        return redirect(url_for('entity_recognition_page'))
    except Exception as e:
        flash(f'添加实体失败: {str(e)}')
        return redirect(url_for('entity_recognition_page'))

# 删除医学实体
@app.route('/delete_entity', methods=['POST'])
def delete_entity():
    """
    删除医学实体的路由
    """
    try:
        entity_type = request.form.get('entity_type')
        entity_name = request.form.get('entity_name')
        
        if not entity_type or not entity_name:
            flash('实体类型和名称不能为空')
            return redirect(url_for('entity_recognition_page'))
        
        # 获取现有实体
        entities = get_medical_entities()
        
        # 如果实体类型存在且实体名称存在，删除实体
        if entity_type in entities and entity_name in entities[entity_type]:
            entities[entity_type].remove(entity_name)
            save_medical_entities(entities)
            flash(f'成功删除实体: {entity_name} (类型: {entity_type})')
        else:
            flash(f'实体不存在: {entity_name} (类型: {entity_type})')
        
        return redirect(url_for('entity_recognition_page'))
    except Exception as e:
        flash(f'删除实体失败: {str(e)}')
        return redirect(url_for('entity_recognition_page'))

# 数据检查页面
@app.route('/check')
def check_page():
    """数据检查页面路由，显示数据检查界面"""
    rules = get_rules()
    return render_template('check.html', rules=rules)

# 清理DataFrame以便JSON序列化
def clean_dataframe_for_json(df):
    """
    清理DataFrame以便JSON序列化，处理特殊值如NaT、NaN等
    
    Args:
        df (DataFrame): 要清理的DataFrame
        
    Returns:
        list: 清理后的数据记录列表
    """
    # 创建DataFrame的副本，避免修改原始数据
    df_clean = df.copy()
    
    # 将NaT（Not a Time）转换为None
    for col in df_clean.select_dtypes(include=['datetime64']).columns:
        df_clean[col] = df_clean[col].astype(object).where(~df_clean[col].isna(), None)
    
    # 将NaN转换为None
    df_clean = df_clean.where(pd.notnull(df_clean), None)
    
    # 转换为字典
    return df_clean.to_dict(orient='records')

# 上传并检查文件
@app.route('/upload', methods=['POST'])
def upload_file():
    """
    处理文件上传和规则检查的路由
    读取上传的Excel文件，执行规则检查，并返回检查结果
    """
    if 'file' not in request.files:
        flash('没有选择文件')
        return redirect(url_for('check_page'))
    
    file = request.files['file']
    if file.filename == '':
        flash('没有选择文件')
        return redirect(url_for('check_page'))
    
    try:
        # 保存上传的文件
        file_path = os.path.join('uploads', file.filename)
        file.save(file_path)
        
        # 读取Excel文件
        if file.filename.endswith('.xlsx') or file.filename.endswith('.xls'):
            df = pd.read_excel(file_path)
        else:
            flash('请上传Excel文件（.xlsx或.xls格式）')
            return redirect(url_for('check_page'))
            
        # 确保DataFrame非空
        if df.empty:
            flash('上传的文件不包含任何数据')
            return redirect(url_for('check_page'))
            
        # 执行规则检查
        results = check_rules(df)
        
        # 处理特殊值
        try:
            # 将NaT（Not a Time）转换为字符串"NaT"
            for col in df.select_dtypes(include=['datetime64']).columns:
                df[col] = df[col].astype(str).replace('NaT', '')
                
            # 确保所有对象类型的列都转换为字符串
            for col in df.select_dtypes(include=['object']).columns:
                df[col] = df[col].astype(str)
            
            # 处理数值列
            for col in df.select_dtypes(include=['number']).columns:
                df[col] = df[col].astype(str)
        except Exception as e:
            print(f"转换数据类型时出错: {str(e)}")
            # 如果出错，尝试将所有列转换为字符串
            for col in df.columns:
                try:
                    df[col] = df[col].astype(str)
                except:
                    pass
        
        # 将NaN值替换为空字符串，避免显示"nan"
        df = df.fillna('')
        df = df.replace('nan', '')
        
        # 转换为HTML表格，不显示行索引
        data_html = df.to_html(classes='table table-striped', index=False)
        
        return render_template('results.html', results=results, data=data_html)
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"文件处理错误: {str(e)}\n{error_details}")
        flash(f'文件处理错误: {str(e)}')
        return redirect(url_for('check_page'))

# 执行规则检查
def check_rules(data):
    """
    根据规则检查数据，找出不符合规则的记录
    
    对输入的数据应用所有质控规则，检查每条记录是否符合规则要求。
    支持缺项检查、逻辑检查和关联逻辑检查三种规则类型。
    
    Args:
        data (pandas.DataFrame): 要检查的数据，每行为一条记录
        
    Returns:
        tuple: (错误列表, 有问题的行索引集合)
            - 错误列表: 包含每个错误的详细信息（规则名称、错误信息、行号等）
            - 有问题的行索引集合: 包含所有存在问题的行的索引
    """
    rules = get_rules()
    results = []
    
    for rule in rules:
        try:
            rule_type = rule['type']
            message = rule['message']
            
            # 缺项检查
            if rule_type == 'missing':
                field = rule['field']
                condition = rule['condition']
                
                # 跳过数据中不存在的字段
                if field not in data.columns:
                    continue
                    
                # 检查字段是否为空或非空
                mask = data[field].isna() | (data[field] == '')
                if condition == 'equals':  # 等于空
                    errors = data[mask]
                else:  # not_equals，不等于空
                    errors = data[~mask]
                    
            # 逻辑检查
            elif rule_type == 'logic':
                field = rule['field']
                condition = rule['condition']
                value = rule['value']
                
                # 跳过数据中不存在的字段
                if field not in data.columns:
                    continue
                    
                # 根据不同的条件类型执行逻辑检查
                if condition == 'equals':  # 等于
                    mask = data[field] == value
                elif condition == 'not_equals':  # 不等于
                    mask = data[field] != value
                elif condition == 'greater_than':  # 大于
                    # 处理日期或数值比较的情况
                    if value in data.columns:  # 如果value是另一个字段名
                        try:
                            # 尝试将两个字段转换为日期类型进行比较
                            field_dates = pd.to_datetime(data[field], errors='coerce')
                            value_dates = pd.to_datetime(data[value], errors='coerce')
                            mask = field_dates <= value_dates  # 检查field是否小于等于value
                        except:
                            # 如果转换失败，尝试直接比较
                            mask = data[field] <= data[value]
                    else:
                        try:
                            # 尝试将字段转换为数值进行比较
                            mask = pd.to_numeric(data[field], errors='coerce') <= float(value)
                        except:
                            # 如果转换失败，尝试字符串比较
                            mask = data[field].astype(str) <= str(value)
                elif condition == 'less_than':  # 小于
                    if value in data.columns:  # 如果value是另一个字段名
                        try:
                            # 尝试将两个字段转换为日期类型进行比较
                            field_dates = pd.to_datetime(data[field], errors='coerce')
                            value_dates = pd.to_datetime(data[value], errors='coerce')
                            mask = field_dates >= value_dates  # 检查field是否大于等于value
                        except:
                            # 如果转换失败，尝试直接比较
                            mask = data[field] >= data[value]
                    else:
                        try:
                            # 尝试将字段转换为数值进行比较
                            mask = pd.to_numeric(data[field], errors='coerce') >= float(value)
                        except:
                            # 如果转换失败，尝试字符串比较
                            mask = data[field].astype(str) >= str(value)
                elif condition == 'contains':  # 包含
                    mask = data[field].astype(str).str.contains(value)
                elif condition == 'not_contains':  # 不包含
                    mask = ~data[field].astype(str).str.contains(value)
                
                # 根据条件名称判断是否需要反转结果
                if 'not' in condition:
                    errors = data[~mask]
                else:
                    errors = data[mask]
                    
            # 关联逻辑检查
            elif rule_type == 'relation':
                field1 = rule['field1']
                field2 = rule['field2']
                relation = rule['relation']
                value_pairs = rule['value_pairs']
                
                # 跳过数据中不存在的字段
                if field1 not in data.columns or field2 not in data.columns:
                    continue
                
                # 特殊处理：科室与入院诊断匹配
                if (field1 == '科室' and field2 == '入院诊断' and relation == 'match_diagnosis') or \
                   (field1 == '科室' and field2 == '主要诊断' and relation == 'match_diagnosis'):
                    try:
                        # 获取科室与诊断映射
                        dept_diag_mapping = get_diagnosis_dept_mapping()
                        errors = pd.DataFrame()
                        
                        # 遍历每行数据
                        for idx, row in data.iterrows():
                            dept = row[field1]
                            diagnosis = row[field2]
                            
                            # 跳过空值
                            if pd.isna(dept) or pd.isna(diagnosis) or dept == '' or diagnosis == '':
                                continue
                            
                            # 如果科室存在于映射中
                            if dept in dept_diag_mapping:
                                # 检查诊断是否与科室匹配
                                matched = False
                                # 遍历科室对应的诊断列表
                                for valid_diag in dept_diag_mapping[dept]:
                                    # 如果诊断包含有效诊断关键词，则认为匹配
                                    if valid_diag in diagnosis:
                                        matched = True
                                        break
                                
                                # 对于特殊的"外科"，需要进一步细分判断
                                if dept.endswith("外科") and not matched:
                                    # 检查是否匹配其他细分外科
                                    for specific_dept in [d for d in dept_diag_mapping if d.endswith("外科") and d != dept]:
                                        for valid_diag in dept_diag_mapping[specific_dept]:
                                            if valid_diag in diagnosis:
                                                matched = True
                                                break
                                        if matched:
                                            break
                                
                                # 对于肿瘤科，特殊处理肺癌、胃癌等多种癌症情况
                                if not matched and "癌" in diagnosis:
                                    for dept_name, diagnoses in dept_diag_mapping.items():
                                        if any(diag in diagnosis for diag in diagnoses if "癌" in diag):
                                            if dept != dept_name and dept_name != "肿瘤科":
                                                matched = False
                                                break
                                
                                # 如果不匹配，添加到错误结果中
                                if not matched:
                                    errors = pd.concat([errors, data.iloc[[idx]]])
                        
                        # 如果有错误，添加到结果中
                        if len(errors) > 0:
                            result = {
                                'rule_name': rule['name'],
                                'message': message,
                                'error_count': len(errors),
                                'error_indices': errors.index.tolist()
                            }
                            results.append(result)
                        continue
                    except Exception as e:
                        # 出错时继续使用常规处理方式
                        print(f"诊断科室匹配检查出错: {str(e)}")
                        pass
                
                # 特殊处理：年龄>14但挂了儿科
                if field1 == '年龄' and field2 == '科室' and relation == 'not_match':
                    try:
                        # 将年龄转换为数值
                        age_numeric = pd.to_numeric(data[field1], errors='coerce')
                        # 检查年龄>14且科室为儿科的情况
                        mask = (age_numeric > 14) & (data[field2].str.contains('儿科'))
                        if mask.any():
                            errors = data[mask]
                            result = {
                                'rule_name': rule['name'],
                                'message': message,
                                'error_count': len(errors),
                                'error_indices': errors.index.tolist()
                            }
                            results.append(result)
                        continue
                    except Exception as e:
                        # 出错时继续使用常规处理方式
                        print(f"年龄科室匹配检查出错: {str(e)}")
                        pass
                
                # 常规关联逻辑检查处理
                try:
                    # 解析值对列表
                    value_pairs_list = []
                    for pair in value_pairs.strip().split('\n'):
                        if pair.strip():
                            val1, val2 = pair.split('=')
                            value_pairs_list.append((val1.strip(), val2.strip()))
                    
                    # 执行匹配检查
                    if relation == 'match':
                        # 检查字段1和字段2的值是否符合指定的匹配关系
                        errors = pd.DataFrame()
                        for val1, val2 in value_pairs_list:
                            # 找出字段1等于val1但字段2不等于val2的行
                            if ',' in val2:  # 如果val2包含多个可能的值
                                valid_val2s = [v.strip() for v in val2.split(',')]
                                mask = (data[field1] == val1) & (~data[field2].isin(valid_val2s))
                            else:
                                mask = (data[field1] == val1) & (data[field2] != val2)
                            errors = pd.concat([errors, data[mask]])
                    elif relation == 'not_match':
                        # 检查字段1和字段2的值是否不符合指定的匹配关系
                        errors = pd.DataFrame()
                        for val1, val2 in value_pairs_list:
                            # 找出字段1等于val1且字段2等于val2的行
                            if ',' in val2:  # 如果val2包含多个不允许的值
                                valid_val2s = [v.strip() for v in val2.split(',')]
                                mask = (data[field1] == val1) & (data[field2].isin(valid_val2s))
                            else:
                                mask = (data[field1] == val1) & (data[field2] == val2)
                            errors = pd.concat([errors, data[mask]])
                except Exception as e:
                    # 解析值对失败时跳过此规则
                    print(f"解析值对失败: {str(e)}")
                    continue
            
            # 跳过没有错误的规则
            if 'errors' not in locals() or len(errors) == 0:
                continue
                
            # 收集错误信息
            error_indices = errors.index.tolist()
            result = {
                'rule_name': rule['name'],
                'message': message,
                'error_count': len(error_indices),
                'error_indices': error_indices
            }
            results.append(result)
            
        except Exception as e:
            # 捕获规则执行过程中的任何异常，确保一个规则的错误不会影响其他规则
            print(f"规则执行出错 ({rule.get('name', '未命名规则')}): {str(e)}")
            continue
    
    return results

# 导出检查结果为CSV的路由
@app.route('/export_results', methods=['POST'])
def export_results():
    """
    导出检查结果的路由
    将检查结果导出为CSV文件供下载
    """
    try:
        # 从POST数据中获取结果和原始数据
        results_data = request.form.get('results')
        original_data = request.form.get('original_data')
        
        if not results_data or not original_data:
            flash('没有可导出的数据')
            return redirect(url_for('check_page'))
            
        # 解析JSON数据
        results = json.loads(results_data)
        original_df = pd.read_json(original_data, orient='records')
        
        # 创建结果文件名（带时间戳）
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        result_filename = f"质控结果_{timestamp}.csv"
        result_path = os.path.join('uploads', result_filename)
        
        # 创建结果DataFrame
        result_df = pd.DataFrame(columns=['规则名称', '错误信息', '错误行号'])
        
        for result in results:
            rule_name = result['rule_name']
            message = result['message']
            for idx in result['error_indices']:
                result_df = pd.concat([result_df, pd.DataFrame({
                    '规则名称': [rule_name],
                    '错误信息': [message],
                    '错误行号': [idx + 1]  # 加1以匹配Excel的行号
                })])
        
        # 保存结果到CSV
        result_df.to_csv(result_path, index=False, encoding='utf-8-sig')
        
        # 返回下载链接
        return jsonify({
            'success': True,
            'filename': result_filename,
            'download_url': url_for('download_file', filename=result_filename)
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

# 文件下载路由
@app.route('/download/<filename>')
def download_file(filename):
    """
    文件下载路由
    
    Args:
        filename (str): 要下载的文件名
    """
    # 检查文件是否在uploads目录中
    if os.path.exists(os.path.join('uploads', filename)):
        return send_from_directory('uploads', filename, as_attachment=True)
    # 检查文件是否在excel_data目录中
    elif os.path.exists(os.path.join('excel_data', filename)):
        return send_from_directory('excel_data', filename, as_attachment=True)
    else:
        flash(f'文件不存在: {filename}')
        return redirect(url_for('index'))

# 添加规则
@app.route('/rules/add', methods=['POST'])
def add_rule():
    """
    添加新规则的路由
    处理表单提交的规则数据，根据规则类型处理不同的字段
    """
    try:
        rules = get_rules()
        rule_id = datetime.now().strftime('%Y%m%d%H%M%S')
        
        rule_type = request.form.get('type')
        
        # 创建基本规则结构
        new_rule = {
            'id': rule_id,
            'name': request.form.get('name'),
            'type': rule_type,
            'message': request.form.get('message')
        }
        
        # 根据规则类型添加不同的字段
        if rule_type == 'missing':  # 缺项检查
            new_rule.update({
                'field': request.form.get('field'),
                'condition': request.form.get('condition'),
                'value': request.form.get('value')
            })
        elif rule_type == 'logic':  # 逻辑检查
            new_rule.update({
                'field': request.form.get('field'),
                'condition': request.form.get('condition'),
                'value': request.form.get('value')
            })
        elif rule_type == 'relation':  # 关联逻辑检查
            new_rule.update({
                'field1': request.form.get('field1'),
                'field2': request.form.get('field2'),
                'relation': request.form.get('relation'),
                'value_pairs': request.form.get('value_pairs')
            })
        
        rules.append(new_rule)
        save_rules(rules)
        flash('规则添加成功！')
    except Exception as e:
        flash(f'添加规则失败: {str(e)}')
    
    return redirect(url_for('rules_page'))

# 编辑规则
@app.route('/rules/edit/<rule_id>', methods=['POST'])
def edit_rule(rule_id):
    """
    编辑现有规则的路由
    
    Args:
        rule_id (str): 要编辑的规则ID
    """
    try:
        rules = get_rules()
        rule_found = False
        
        for rule in rules:
            if rule['id'] == rule_id:
                rule_found = True
                rule_type = request.form.get('type')
                
                # 更新基本信息
                rule['name'] = request.form.get('name')
                rule['type'] = rule_type
                rule['message'] = request.form.get('message')
                
                # 根据规则类型更新不同的字段
                if rule_type == 'missing':  # 缺项检查
                    rule['field'] = request.form.get('field')
                    rule['condition'] = request.form.get('condition')
                    rule['value'] = request.form.get('value')
                elif rule_type == 'logic':  # 逻辑检查
                    rule['field'] = request.form.get('field')
                    rule['condition'] = request.form.get('condition')
                    rule['value'] = request.form.get('value')
                elif rule_type == 'relation':  # 关联逻辑检查
                    rule['field1'] = request.form.get('field1')
                    rule['field2'] = request.form.get('field2')
                    rule['relation'] = request.form.get('relation')
                    rule['value_pairs'] = request.form.get('value_pairs')
                
                break
        
        if not rule_found:
            flash('规则不存在！')
            return redirect(url_for('rules_page'))
            
        save_rules(rules)
        flash('规则更新成功！')
    except Exception as e:
        flash(f'更新规则失败: {str(e)}')
    
    return redirect(url_for('rules_page'))

# 删除规则
@app.route('/rules/delete/<rule_id>')
def delete_rule(rule_id):
    """
    删除规则的路由
    
    Args:
        rule_id (str): 要删除的规则ID
    """
    try:
        rules = get_rules()
        rules = [rule for rule in rules if rule['id'] != rule_id]
        save_rules(rules)
        flash('规则删除成功！')
    except Exception as e:
        flash(f'删除规则失败: {str(e)}')
    
    return redirect(url_for('rules_page'))

# 大模型命名实体识别页面
@app.route('/llm_entity_recognition')
def llm_entity_recognition_page():
    """大模型命名实体识别页面路由，显示大模型实体识别界面"""
    llm_config = get_llm_config()
    
    # 添加提示信息
    messages = []
    
    # 检查是否安装了requests库（用于API调用）
    try:
        import requests
        api_available = True
    except ImportError:
        api_available = False
        messages.append({
            'type': 'warning',
            'content': '未检测到requests库，系统无法使用API模式。如需使用API模式，请安装requests库。'
        })
    
    # 检查是否处于API模式
    if llm_config.get("api_mode", False) and api_available:
        messages.append({
            'type': 'info',
            'content': f'系统当前处于API模式，将使用{llm_config.get("api_type", "未指定")} API进行实体识别。'
        })
    else:
        messages.append({
            'type': 'warning',
            'content': '系统当前未启用API模式，将使用规则匹配进行实体识别。请在配置中启用API模式获得更好的识别效果。'
        })
    
    return render_template('llm_entity_recognition.html', 
                           llm_config=llm_config, 
                           api_available=api_available,
                           messages=messages)

# 处理大模型命名实体识别请求
@app.route('/recognize_llm_entities', methods=['POST'])
def recognize_llm_entities():
    """
    处理基于大模型的医学实体识别请求的路由
    
    接收用户提交的医学文本，使用API模式或规则匹配方法识别其中的医学实体，
    并返回高亮显示识别结果的页面，同时生成实体类型分布的饼状图。
    支持文本输入和文件上传两种方式。
    
    Returns:
        str: 渲染后的大模型实体识别结果页面HTML，包含实体统计图表
    """
    try:
        # 获取表单数据
        text = request.form.get('text', '')
        
        # 处理文件上传
        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']
            if file.filename.endswith('.txt'):
                try:
                    text = file.read().decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        # 尝试其他编码
                        text = file.read().decode('gbk')
                    except:
                        flash('无法解析文件编码，请使用UTF-8或GBK编码的文本文件')
                        return redirect(url_for('llm_entity_recognition_page'))
            else:
                flash('仅支持.txt文本文件')
                return redirect(url_for('llm_entity_recognition_page'))
        
        if not text:
            flash('请输入文本或上传文本文件')
            return redirect(url_for('llm_entity_recognition_page'))
        
        # 获取配置
        config = get_llm_config()
        api_mode = config.get("api_mode", False)
        
        # 根据选择的模型类型进行实体识别
        if api_mode:
            # API模式下使用API识别
            recognized_entities = recognize_entities_with_api(text)
            model_type = f"{config.get('api_type', 'API')}模式"
        else:
            # 使用规则匹配
            recognized_entities = recognize_entities_with_rules(text)
            model_type = "规则匹配"
        
        if not recognized_entities:
            flash('未能识别到任何实体，请检查模型配置或尝试其他文本')
            return redirect(url_for('llm_entity_recognition_page'))
        
        # 统计实体频率
        entity_statistics = calculate_entity_statistics(recognized_entities)
        
        # 准备高亮显示的文本
        highlighted_text = text
        
        # 按位置排序所有实体（从后往前替换，避免位置偏移）
        all_entities = []
        for entity_type, entities in recognized_entities.items():
            for entity_info in entities:
                all_entities.append({
                    'entity': entity_info['entity'],
                    'position': entity_info['position'],
                    'type': entity_type
                })
        
        # 按位置从大到小排序
        all_entities.sort(key=lambda x: x['position'], reverse=True)
        
        # 替换文本中的实体为高亮标签
        for entity_info in all_entities:
            entity = entity_info['entity']
            pos = entity_info['position']
            entity_type = entity_info['type']
            
            # 确保位置有效
            if pos >= 0 and pos + len(entity) <= len(highlighted_text):
                highlighted_text = (
                    highlighted_text[:pos] + 
                    f'<span class="entity-highlight {entity_type.lower()}" title="{entity_type}">{entity}</span>' + 
                    highlighted_text[pos+len(entity):]
                )
        
        return render_template('llm_entity_recognition_result.html', 
                               original_text=text,
                               highlighted_text=highlighted_text,
                               recognized_entities=recognized_entities,
                               entity_statistics=entity_statistics,
                               model_type=model_type)
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"大模型实体识别错误: {str(e)}\n{error_details}")
        flash(f'大模型实体识别错误: {str(e)}')
        return redirect(url_for('llm_entity_recognition_page'))

# 保存大模型配置
@app.route('/save_llm_config_route', methods=['POST'])
def save_llm_config_route():
    """
    保存大模型配置的路由
    接收表单数据并更新配置文件
    """
    try:
        # 获取当前配置
        config = get_llm_config()
        
        # 更新配置，只保留API相关配置
        config["api_mode"] = 'api_mode' in request.form
        config["api_type"] = request.form.get('api_type', 'deepseek')
        config["api_key"] = request.form.get('api_key', '')
        config["api_url"] = request.form.get('api_url', '')
        
        # 保存配置
        save_llm_config(config)
        
        flash('配置已保存')
    except Exception as e:
        flash(f'保存配置出错: {str(e)}')
    
    return redirect(url_for('llm_entity_recognition_page'))

# Excel数据检查页面
@app.route('/excel_check')
def excel_check_page():
    """Excel数据检查页面路由，显示Excel数据检查界面"""
    rules = get_rules()
    return render_template('excel_check.html', rules=rules)

# Word文档数据检查页面
@app.route('/docx_check')
def docx_check_page():
    """Word文档数据检查页面路由，显示Word文档数据检查界面"""
    rules = get_rules()
    return render_template('docx_check.html', rules=rules)

# 上传并检查Word文档
@app.route('/upload_docx', methods=['POST'])
def upload_docx():
    """
    处理Word文档上传和规则检查的路由
    读取上传的Word文档，提取数据，执行规则检查，并返回检查结果
    """
    if 'file' not in request.files:
        flash('没有选择文件')
        return redirect(url_for('docx_check_page'))
    
    file = request.files['file']
    if file.filename == '':
        flash('没有选择文件')
        return redirect(url_for('docx_check_page'))
    
    try:
        # 保存上传的文件
        file_path = os.path.join('uploads', secure_filename(file.filename))
        file.save(file_path)
        
        # 检查文件类型
        if not file.filename.endswith('.docx'):
            flash('请上传Word文档(.docx格式)')
            return redirect(url_for('docx_check_page'))
            
        # 提取数据
        extractor = DocxDataExtractor()
        df = extractor.extract_and_convert(file_path)
        
        # 确保DataFrame非空
        if df.empty:
            flash('无法从文档中提取有效数据')
            return redirect(url_for('docx_check_page'))
            
        # 执行规则检查
        results = check_rules(df)
        
        # 获取错误字段列表和规则类型映射
        error_fields = set()
        field_rule_types = {}  # 字段对应的规则类型
        field_related_fields = {}  # 关联字段（用于关键逻辑检查）
        
        # 调试信息
        print(f"检测到 {len(results)} 个错误")
        
        if results:
            for error in results:
                if 'rule_name' in error and 'message' in error:
                    rule_name = error['rule_name']
                    message = error['message']
                    
                    # 调试信息
                    print(f"规则名称: {rule_name}, 错误消息: {message}")
                    
                    # 从规则中提取类型和字段信息
                    rules = get_rules()
                    rule_info = None
                    for r in rules:
                        if r['name'] == rule_name:
                            rule_info = r
                            break
                    
                    if rule_info:
                        rule_type = rule_info['type']
                        
                        # 缺项检查
                        if rule_type == 'missing':
                            field = rule_info['field']
                            error_fields.add(field)
                            field_rule_types[field] = 'missing'
                            print(f"缺项检查字段: {field}")
                        
                        # 逻辑检查
                        elif rule_type == 'logic':
                            field = rule_info['field']
                            error_fields.add(field)
                            field_rule_types[field] = 'logic'
                            print(f"逻辑检查字段: {field}")
                        
                        # 关联逻辑检查
                        elif rule_type == 'relation':
                            field1 = rule_info['field1']
                            field2 = rule_info['field2']
                            error_fields.add(field1)
                            error_fields.add(field2)
                            field_rule_types[field1] = 'relation'
                            field_rule_types[field2] = 'relation'
                            field_related_fields[field1] = field2
                            field_related_fields[field2] = field1
                            print(f"关联逻辑检查字段: {field1}, {field2}")
                    
                    # 如果没有找到规则信息，尝试从规则名称中提取字段
                    else:
                        if '缺' in rule_name or '漏' in rule_name:
                            field_match = re.search(r'([^缺漏]+)[缺漏]', rule_name)
                            if field_match:
                                field = field_match.group(1).strip()
                                error_fields.add(field)
                                field_rule_types[field] = 'missing'  # 假设是缺项检查
                                print(f"从规则名称提取字段(缺项): {field}")
        
        # 调试信息
        print(f"最终错误字段列表: {error_fields}")
        print(f"字段规则类型映射: {field_rule_types}")
        print(f"关联字段映射: {field_related_fields}")
        
        # 生成标记错误的Word文档
        result_docx_path = None
        if results:
            result_generator = DocxResultGenerator()
            result_docx_path = result_generator.highlight_errors(file_path, results)
            if result_docx_path:
                result_docx_path = os.path.basename(result_docx_path)
        
        # 处理特殊值
        try:
            # 将NaT（Not a Time）转换为字符串"NaT"
            for col in df.select_dtypes(include=['datetime64']).columns:
                df[col] = df[col].astype(str).replace('NaT', '')
                
            # 确保所有对象类型的列都转换为字符串
            for col in df.select_dtypes(include=['object']).columns:
                df[col] = df[col].astype(str)
            
            # 处理数值列
            for col in df.select_dtypes(include=['number']).columns:
                df[col] = df[col].astype(str)
        except Exception as e:
            print(f"转换数据类型时出错: {str(e)}")
            # 如果出错，尝试将所有列转换为字符串
            for col in df.columns:
                try:
                    df[col] = df[col].astype(str)
                except:
                    pass
        
        # 将NaN值替换为空字符串，避免显示"nan"
        df = df.fillna('')
        df = df.replace('nan', '')
        
        # 转换为HTML表格，不显示行索引
        data_html = df.to_html(classes='table table-striped', index=False)
        
        # 提取文档内容用于HTML预览，并传递错误字段信息和规则类型
        docx_html = extract_docx_html(file_path, error_fields, field_rule_types, field_related_fields)
        
        return render_template('docx_results.html', 
                              results=results, 
                              data=data_html, 
                              docx_html=docx_html,
                              result_docx_path=result_docx_path)
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Word文档处理错误: {str(e)}\n{error_details}")
        flash(f'Word文档处理错误: {str(e)}')
        return redirect(url_for('docx_check_page'))

# 提取Word文档内容为HTML预览
def extract_docx_html(docx_path, error_fields, field_rule_types, field_related_fields):
    """
    提取Word文档内容用于HTML预览
    
    Args:
        docx_path (str): Word文档路径
        error_fields (set): 错误字段集合
        field_rule_types (dict): 字段对应的规则类型
        field_related_fields (dict): 关联字段映射
        
    Returns:
        str: HTML格式的文档内容
    """
    try:
        from docx import Document
        from docx.document import Document as DocumentType
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        doc = Document(docx_path)
        html_parts = ['<div class="docx-content">']
        
        # 调试信息
        print(f"错误字段集合: {error_fields}")
        print(f"字段规则类型映射: {field_rule_types}")
        
        # 简化的字段匹配模式，确保能够匹配到字段
        field_patterns = {}
        for field in error_fields:
            # 创建一个简单的模式，匹配字段名后面跟着冒号、下划线或空格
            field_patterns[field] = re.compile(f"{re.escape(field)}[：:_\\s]*")
        
        # 获取文档中所有元素（段落和表格）并按顺序处理
        def process_paragraph(para):
            if para.text.strip():
                # 检查段落是否包含错误字段
                para_text = para.text
                marked_text = html.escape(para_text)
                
                for field, pattern in field_patterns.items():
                    match = pattern.search(para_text)
                    if match:
                        # 获取字段在文本中的位置
                        start_pos = match.start()
                        field_end = match.end()
                        
                        # 根据规则类型决定如何标记
                        rule_type = field_rule_types.get(field, 'missing')  # 默认为缺项检查
                        
                        if rule_type == 'missing':
                            # 缺项检查：只标记字段
                            marked_text = (
                                html.escape(para_text[:start_pos]) +
                                f'<span class="text-danger fw-bold">{html.escape(para_text[start_pos:field_end])}</span>' +
                                html.escape(para_text[field_end:])
                            )
                        
                        elif rule_type == 'logic':
                            # 逻辑检查：标记字段和值
                            # 查找字段对应的值区域（从字段结束到下一个字段开始或行尾）
                            value_end = len(para_text)
                            for next_field, next_pattern in field_patterns.items():
                                if next_field != field:
                                    next_match = next_pattern.search(para_text[field_end:])
                                    if next_match:
                                        value_end = field_end + next_match.start()
                                        break
                            
                            # 标记字段和值
                            marked_text = (
                                html.escape(para_text[:start_pos]) +
                                f'<span class="text-danger fw-bold">{html.escape(para_text[start_pos:value_end])}</span>' +
                                html.escape(para_text[value_end:])
                            )
                        
                        elif rule_type == 'relation':
                            # 关联逻辑检查：标记字段1和字段2
                            related_field = field_related_fields.get(field)
                            if related_field:
                                # 标记当前字段
                                marked_text = (
                                    html.escape(para_text[:start_pos]) +
                                    f'<span class="text-danger fw-bold">{html.escape(para_text[start_pos:field_end])}</span>' +
                                    html.escape(para_text[field_end:])
                                )
                                
                                # 检查是否同时包含关联字段
                                related_pattern = field_patterns.get(related_field)
                                if related_pattern:
                                    related_match = related_pattern.search(para_text)
                                    if related_match:
                                        related_start = related_match.start()
                                        related_end = related_match.end()
                                        
                                        # 重新构建标记文本，同时标记两个字段
                                        if related_start < start_pos:
                                            # 关联字段在前
                                            marked_text = (
                                                html.escape(para_text[:related_start]) +
                                                f'<span class="text-danger fw-bold">{html.escape(para_text[related_start:related_end])}</span>' +
                                                html.escape(para_text[related_end:start_pos]) +
                                                f'<span class="text-danger fw-bold">{html.escape(para_text[start_pos:field_end])}</span>' +
                                                html.escape(para_text[field_end:])
                                            )
                                        else:
                                            # 关联字段在后
                                            marked_text = (
                                                html.escape(para_text[:start_pos]) +
                                                f'<span class="text-danger fw-bold">{html.escape(para_text[start_pos:field_end])}</span>' +
                                                html.escape(para_text[field_end:related_start]) +
                                                f'<span class="text-danger fw-bold">{html.escape(para_text[related_start:related_end])}</span>' +
                                                html.escape(para_text[related_end:])
                                            )
                        
                        break  # 一个段落只处理一个错误字段
                
                return f'<p>{marked_text}</p>'
            return ''
        
        def process_table(table):
            table_html = ['<table class="table table-bordered table-sm">']
            for row in table.rows:
                table_html.append('<tr>')
                for cell in row.cells:
                    cell_text = ' '.join([p.text for p in cell.paragraphs])
                    
                    # 检查单元格是否包含错误字段
                    marked_text = html.escape(cell_text)
                    
                    for field, pattern in field_patterns.items():
                        match = pattern.search(cell_text)
                        if match:
                            # 获取字段在文本中的位置
                            start_pos = match.start()
                            field_end = match.end()
                            
                            # 根据规则类型决定如何标记
                            rule_type = field_rule_types.get(field, 'missing')  # 默认为缺项检查
                            
                            if rule_type == 'missing':
                                # 缺项检查：只标记字段
                                marked_text = (
                                    html.escape(cell_text[:start_pos]) +
                                    f'<span class="text-danger fw-bold">{html.escape(cell_text[start_pos:field_end])}</span>' +
                                    html.escape(cell_text[field_end:])
                                )
                            
                            elif rule_type == 'logic':
                                # 逻辑检查：标记字段和值
                                # 查找字段对应的值区域（从字段结束到下一个字段开始或行尾）
                                value_end = len(cell_text)
                                for next_field, next_pattern in field_patterns.items():
                                    if next_field != field:
                                        next_match = next_pattern.search(cell_text[field_end:])
                                        if next_match:
                                            value_end = field_end + next_match.start()
                                            break
                                
                                # 标记字段和值
                                marked_text = (
                                    html.escape(cell_text[:start_pos]) +
                                    f'<span class="text-danger fw-bold">{html.escape(cell_text[start_pos:value_end])}</span>' +
                                    html.escape(cell_text[value_end:])
                                )
                            
                            elif rule_type == 'relation':
                                # 关联逻辑检查：标记字段1和字段2
                                related_field = field_related_fields.get(field)
                                if related_field:
                                    # 标记当前字段
                                    marked_text = (
                                        html.escape(cell_text[:start_pos]) +
                                        f'<span class="text-danger fw-bold">{html.escape(cell_text[start_pos:field_end])}</span>' +
                                        html.escape(cell_text[field_end:])
                                    )
                                    
                                    # 检查是否同时包含关联字段
                                    related_pattern = field_patterns.get(related_field)
                                    if related_pattern:
                                        related_match = related_pattern.search(cell_text)
                                        if related_match:
                                            related_start = related_match.start()
                                            related_end = related_match.end()
                                            
                                            # 重新构建标记文本，同时标记两个字段
                                            if related_start < start_pos:
                                                # 关联字段在前
                                                marked_text = (
                                                    html.escape(cell_text[:related_start]) +
                                                    f'<span class="text-danger fw-bold">{html.escape(cell_text[related_start:related_end])}</span>' +
                                                    html.escape(cell_text[related_end:start_pos]) +
                                                    f'<span class="text-danger fw-bold">{html.escape(cell_text[start_pos:field_end])}</span>' +
                                                    html.escape(cell_text[field_end:])
                                                )
                                            else:
                                                # 关联字段在后
                                                marked_text = (
                                                    html.escape(cell_text[:start_pos]) +
                                                    f'<span class="text-danger fw-bold">{html.escape(cell_text[start_pos:field_end])}</span>' +
                                                    html.escape(cell_text[field_end:related_start]) +
                                                    f'<span class="text-danger fw-bold">{html.escape(cell_text[related_start:related_end])}</span>' +
                                                    html.escape(cell_text[related_end:])
                                                )
                            
                            break  # 一个单元格只处理一个错误字段
                    
                    table_html.append(f'<td>{marked_text}</td>')
                table_html.append('</tr>')
            table_html.append('</table>')
            return '\n'.join(table_html)
        
        # 按顺序遍历文档中的所有元素
        # python-docx不直接提供按顺序访问所有元素的方法，但我们可以通过文档的_element属性来获取
        try:
            # 尝试获取文档中所有块级元素的顺序
            all_elements = []
            body = doc._element.body
            for child in body.iterchildren():
                if child.tag.endswith('p'):  # 段落
                    idx = [i for i, p in enumerate(doc.paragraphs) if p._element is child]
                    if idx:
                        all_elements.append(('paragraph', doc.paragraphs[idx[0]]))
                elif child.tag.endswith('tbl'):  # 表格
                    idx = [i for i, t in enumerate(doc.tables) if t._element is child]
                    if idx:
                        all_elements.append(('table', doc.tables[idx[0]]))
            
            # 按顺序处理元素
            for elem_type, element in all_elements:
                if elem_type == 'paragraph':
                    html_part = process_paragraph(element)
                    if html_part:
                        html_parts.append(html_part)
                elif elem_type == 'table':
                    html_parts.append(process_table(element))
        except Exception as e:
            print(f"按顺序处理文档元素时出错: {str(e)}")
            # 回退到分开处理段落和表格
            for para in doc.paragraphs:
                html_part = process_paragraph(para)
                if html_part:
                    html_parts.append(html_part)
            
            for table in doc.tables:
                html_parts.append(process_table(table))
        
        html_parts.append('</div>')
        return '\n'.join(html_parts)
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"提取Word文档HTML预览时出错: {str(e)}\n{error_details}")
        return f'<div class="alert alert-danger">无法提取文档内容: {str(e)}</div>'

# 应用入口
if __name__ == '__main__':
    app.run(debug=True)  # 生产环境应设置debug=False 